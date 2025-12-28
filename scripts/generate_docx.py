#!/usr/bin/env python3
"""
Script de generation de documents DOCX a partir de Markdown.
Utilise python-docx pour creer des documents Word professionnels.

Usage:
    python generate_docx.py <fichier.md> [output.docx]

Exemple:
    python generate_docx.py rapport.md
    python generate_docx.py rapport.md mon_rapport.docx
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Erreur: python-docx n'est pas installe.")
    print("Installation: pip install python-docx")
    sys.exit(1)


def set_cell_shading(cell, color):
    """Applique une couleur de fond a une cellule de tableau."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def parse_markdown(md_content):
    """Parse le contenu Markdown en elements structures."""
    elements = []
    lines = md_content.split('\n')
    current_paragraph = []
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []

    for line in lines:
        # Blocs de code
        if line.strip().startswith('```'):
            if in_code_block:
                elements.append(('code', '\n'.join(code_content)))
                code_content = []
                in_code_block = False
            else:
                if current_paragraph:
                    elements.append(('paragraph', '\n'.join(current_paragraph)))
                    current_paragraph = []
                in_code_block = True
            continue

        if in_code_block:
            code_content.append(line)
            continue

        # Tableaux
        if line.strip().startswith('|') and '|' in line[1:]:
            if not in_table:
                if current_paragraph:
                    elements.append(('paragraph', '\n'.join(current_paragraph)))
                    current_paragraph = []
                in_table = True
            # Ignorer les lignes de separation (|---|---|)
            if not re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                cells = [c.strip() for c in line.strip().split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
            continue
        elif in_table:
            elements.append(('table', table_rows))
            table_rows = []
            in_table = False

        # Ligne horizontale
        if re.match(r'^---+$', line.strip()):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('hr', None))
            continue

        # Titres
        if line.startswith('# '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('heading1', line[2:].strip()))
        elif line.startswith('## '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('heading2', line[3:].strip()))
        elif line.startswith('### '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('heading3', line[4:].strip()))
        elif line.startswith('#### '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('heading4', line[5:].strip()))
        # Citations
        elif line.strip().startswith('> '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('quote', line.strip()[2:]))
        # Listes a puces
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            content = line.strip()[2:]
            # Checkbox
            if content.startswith('[ ] '):
                elements.append(('checkbox_unchecked', content[4:]))
            elif content.startswith('[x] ') or content.startswith('[X] '):
                elements.append(('checkbox_checked', content[4:]))
            else:
                elements.append(('bullet', content))
        # Listes numerotees
        elif re.match(r'^\d+\. ', line.strip()):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('numbered', re.sub(r'^\d+\. ', '', line.strip())))
        # Paragraphes
        elif line.strip():
            current_paragraph.append(line)
        else:
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []

    # Finaliser
    if current_paragraph:
        elements.append(('paragraph', '\n'.join(current_paragraph)))
    if in_table and table_rows:
        elements.append(('table', table_rows))

    return elements


def clean_markdown_text(text):
    """Nettoie le texte Markdown (gras, italique, liens)."""
    # Liens [texte](url) -> texte
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Gras **texte** -> texte
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Italique *texte* -> texte
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Code inline `texte` -> texte
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def add_formatted_paragraph(doc, text, style=None):
    """Ajoute un paragraphe avec formatage basique."""
    clean_text = clean_markdown_text(text)
    p = doc.add_paragraph(clean_text, style=style)
    return p


def create_docx(md_path, output_path=None):
    """Cree un document DOCX a partir d'un fichier Markdown."""
    md_path = Path(md_path)

    if not md_path.exists():
        print(f"Erreur: Le fichier {md_path} n'existe pas.")
        return None

    md_content = md_path.read_text(encoding='utf-8')
    elements = parse_markdown(md_content)

    doc = Document()

    # Configuration des styles
    styles = doc.styles

    # Style pour le code
    if 'Code' not in [s.name for s in styles]:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)

    # Parcourir les elements
    for element_type, content in elements:
        if element_type == 'heading1':
            doc.add_heading(content, level=1)
        elif element_type == 'heading2':
            doc.add_heading(content, level=2)
        elif element_type == 'heading3':
            doc.add_heading(content, level=3)
        elif element_type == 'heading4':
            doc.add_heading(content, level=4)
        elif element_type == 'bullet':
            add_formatted_paragraph(doc, content, style='List Bullet')
        elif element_type == 'numbered':
            add_formatted_paragraph(doc, content, style='List Number')
        elif element_type == 'checkbox_unchecked':
            p = doc.add_paragraph(style='List Bullet')
            p.add_run('☐ ' + clean_markdown_text(content))
        elif element_type == 'checkbox_checked':
            p = doc.add_paragraph(style='List Bullet')
            p.add_run('☑ ' + clean_markdown_text(content))
        elif element_type == 'quote':
            p = add_formatted_paragraph(doc, content)
            p.paragraph_format.left_indent = Cm(1)
            p.runs[0].italic = True
        elif element_type == 'code':
            p = doc.add_paragraph(content, style='Code' if 'Code' in [s.name for s in styles] else None)
            if p.runs:
                p.runs[0].font.name = 'Consolas'
                p.runs[0].font.size = Pt(9)
        elif element_type == 'table':
            if content and len(content) > 0:
                # Creer le tableau
                table = doc.add_table(rows=len(content), cols=len(content[0]))
                table.style = 'Table Grid'

                for i, row_data in enumerate(content):
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        if j < len(row.cells):
                            row.cells[j].text = clean_markdown_text(cell_text)
                            # En-tete en gras
                            if i == 0:
                                for paragraph in row.cells[j].paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                set_cell_shading(row.cells[j], 'E0E0E0')

                doc.add_paragraph()  # Espace apres tableau
        elif element_type == 'hr':
            # Ligne horizontale simulee
            p = doc.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element_type == 'paragraph':
            if content.strip():
                add_formatted_paragraph(doc, content)

    # Sauvegarder
    if output_path is None:
        output_path = str(md_path.with_suffix('.docx'))

    doc.save(output_path)
    print(f"Document DOCX cree : {output_path}")
    return output_path


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    result = create_docx(md_file, output_file)
    if result is None:
        sys.exit(1)


if __name__ == '__main__':
    main()
